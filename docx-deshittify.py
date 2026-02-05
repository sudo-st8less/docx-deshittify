#!/usr/bin/env python3
##############################
##############################
#### docx-deshittify:
#### Removes phantom blank pages, table format metadata, paragraph and space formatting bugs from .docx files.
#### x: @st8less
##############################
##############################

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
import sys
import os


def fix_document_formatting(input_file, output_file):

    print(f"Loading: {input_file}")
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found")
        sys.exit(1)
    
    doc = Document(input_file)
    
    print(f"Found {len(doc.paragraphs)} paragraphs")
    print(f"Found {len(doc.tables)} tables")
    
# squash leading empty paragraphs
    print("\nRemoving leading nullll paragraphs...")
    removed_leading = 0
    while doc.paragraphs:
        para = doc.paragraphs[0]
        if len(para.text.strip()) == 0 and len(para.runs) <= 1:
            if len(para.runs) == 1 and len(para.runs[0].text.strip()) == 0:
                para._element.getparent().remove(para._element)
                removed_leading += 1
            else:
                break
        else:
            break
    print(f"  Removed {removed_leading} leading paragraphs")
    
# remove trailing empty paras
    print("\nRemoving trailing empty paragraphs...")
    removed_trailing = 0
    while doc.paragraphs:
        para = doc.paragraphs[-1]
        if len(para.text.strip()) == 0 and len(para.runs) <= 1:
            if len(para.runs) == 1 and len(para.runs[0].text.strip()) == 0:
                para._element.getparent().remove(para._element)
                removed_trailing += 1
            else:
                break
        else:
            break
    print(f"  Removed {removed_trailing} trailing paragraphs")
    
# spacing issue, created trail blank page on g docs
    print("\nFixing paragraph spacing...")
    spacing_fixed = 0
    for para in doc.paragraphs:
# xml check
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            spacing_elem = pPr.find(qn('w:spacing'))
            if spacing_elem is not None:
                before_val = spacing_elem.get(qn('w:before'))
                after_val = spacing_elem.get(qn('w:after'))
                if before_val or after_val:
                    print(f"  Found spacing - before: {before_val}, after: {after_val}")
                    spacing_fixed += 1
        
# copied spacing error fix, throws excep without
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
    print(f"  Fixed spacing on {spacing_fixed} paragraphs")
    
#table block
    print("\nFixing table formatting...")
    for i, table in enumerate(doc.tables):
        print(f"  Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
        
# allow automatic height
        for row in table.rows:
            row.height = None
            trPr = row._element.find(qn('w:trPr'))
            if trPr is not None:
# Remove cantSplit if it exists
                cantSplit = trPr.find(qn('w:cantSplit'))
                if cantSplit is not None:
                    trPr.remove(cantSplit)
                    print(f"    Removed cantSplit from row")
        
#table cell spacing
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            spacing = tblPr.find(qn('w:tblCellSpacing'))
            if spacing is not None:
                print(f"    Removing table cell spacing")
                tblPr.remove(spacing)
    
# iter over page break data
    print("\nRemoving page breaks...")
    breaks_removed = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for elem in run._element:
                if elem.tag.endswith('br'):
                    br_type = elem.get(qn('w:type'))
                    if br_type == 'page':
                        print(f"  Found and removing page break")
                        elem.getparent().remove(elem)
                        breaks_removed += 1
        
#pageBreakBefore saved prop
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            pageBreakBefore = pPr.find(qn('w:pageBreakBefore'))
            if pageBreakBefore is not None:
                print(f"  Removing pageBreakBefore property")
                pPr.remove(pageBreakBefore)
                breaks_removed += 1
    print(f"  Removed {breaks_removed} page breaks")
    
# margins, works
    print("\nSetting standard margins...")
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    print(f"\nSaving to: {output_file}")
    doc.save(output_file)
    print("Done!")
  
#cli output pretty
    print(f"\nFinal document stats:")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")
    print(f"  Sections: {len(doc.sections)}")
  
def print_usage():
    """Print usage information"""
    print("docx-deshittify - Remove formatting issues from .docx files")
    print("\nUsage:")
    print("  python3 docx-deshittify.py <input.docx> <output.docx>")
  
#main 
if __name__ == "__main__":
    if len(sys.argv) < 2 or sys.argv[1] in ['-h', '--help', 'help']:
        print_usage()
        sys.exit(0)
    
    input_file = sys.argv[1]
    
# output filename 
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
    else:
        input_dir = os.path.dirname(input_file)
        input_basename = os.path.basename(input_file)
        
# prefix for jerks
        if input_dir:
            output_file = os.path.join(input_dir, 'fixed_' + input_basename)
        else:
            output_file = 'fixed_' + input_basename
    
    fix_document_formatting(input_file, output_file)
